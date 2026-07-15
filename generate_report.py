"""Generate the HCI report as a .docx file."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUT = os.path.join(os.path.dirname(__file__), "HCI_Report_Deadline_Extractor.docx")

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1.1)
    section.right_margin  = Inches(1.1)

# ── Style helpers ─────────────────────────────────────────────────────────────
BODY_FONT  = "Calibri"
HEAD_FONT  = "Calibri"
BODY_SIZE  = Pt(11)
HEAD_COLOR = RGBColor(0x1A, 0x1A, 0x2E)   # deep navy


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.name  = HEAD_FONT
    run.font.color.rgb = HEAD_COLOR
    run.font.bold  = True
    run.font.size  = Pt(13) if level == 1 else Pt(11.5)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    return p


def body(doc, text, bold=False, italic=False, indent=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name   = BODY_FONT
    run.font.size   = BODY_SIZE
    run.font.bold   = bold
    run.font.italic = italic
    p.paragraph_format.space_after  = Pt(5)
    p.paragraph_format.space_before = Pt(0)
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.name = BODY_FONT
    run.font.size = BODY_SIZE
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.left_indent  = Inches(0.3 + level * 0.2)
    return p


def add_table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    # header row
    hdr_cells = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.font.bold = True
        run.font.size = Pt(9.5)
        run.font.name = BODY_FONT
        tc = hdr_cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "D9E2F3")
        tcPr.append(shd)
    # data rows
    for ri, row in enumerate(rows):
        cells = t.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
            run = cells[ci].paragraphs[0].runs[0] if cells[ci].paragraphs[0].runs else cells[ci].paragraphs[0].add_run(val)
            run.font.size = Pt(9.5)
            run.font.name = BODY_FONT
    doc.add_paragraph()
    return t


# ═══════════════════════════════════════════════════════════════════════════════
#  TITLE BLOCK
# ═══════════════════════════════════════════════════════════════════════════════
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title_p.add_run("Deadline Extractor")
tr.font.name  = HEAD_FONT
tr.font.size  = Pt(22)
tr.font.bold  = True
tr.font.color.rgb = HEAD_COLOR

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub_p.add_run(
    "Model: Gemma 4 e2b (5.1B, Q4_K_M)  |  "
    "Machine: AMD Ryzen 7 7435HS, 31.8 GB RAM, Windows 11 Home  |  "
    "Approach: Pipeline"
)
sr.font.name   = BODY_FONT
sr.font.size   = Pt(9.5)
sr.font.italic = True

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
#  1. PROBLEM STATEMENT
# ═══════════════════════════════════════════════════════════════════════════════
heading(doc, "Problem Statement")
body(doc,
    "University students and working professionals deal with important deadline information "
    "buried inside long documents every single day. A typical university student gets handed "
    "four or five course syllabi, each one a dense PDF with assignment due dates, exam dates, "
    "lab submission windows, and participation deadlines scattered across ten to thirty pages. "
    "They also receive mid-semester emails from professors pushing dates around. Manually "
    "reading through all of that and copying every date into a calendar is slow, boring, and "
    "dangerously easy to get wrong. Miss one date and an assignment gets a zero."
)
body(doc,
    "The scoped problem this project addresses: given any combination of a PDF document, "
    "an image or screenshot, or raw pasted text, automatically identify every deadline or "
    "time-sensitive event, label it with a human-readable title, and convert the full list "
    "into a downloadable .ics calendar file that imports directly into Google Calendar, "
    "Apple Calendar, or Outlook without any manual re-entry."
)
body(doc,
    "This is narrowed deliberately. The system does not summarise documents, answer questions "
    "about them, or manage tasks. It does one thing: extract dates and hand the user a "
    "calendar file."
)

# ═══════════════════════════════════════════════════════════════════════════════
#  2. PERSONA  (NNGroup template)
# ═══════════════════════════════════════════════════════════════════════════════
heading(doc, "Persona")

# ── Name + tagline block ──────────────────────────────────────────────────────
name_p = doc.add_paragraph()
name_r = name_p.add_run("Maya Chen — ")
name_r.font.name = HEAD_FONT
name_r.font.size = Pt(13)
name_r.font.bold = True
name_r.font.color.rgb = HEAD_COLOR
tag_r = name_p.add_run('"I know all my deadlines are in those PDFs somewhere. '
                        'I just never have time to dig them all out."')
tag_r.font.name   = BODY_FONT
tag_r.font.size   = Pt(11)
tag_r.font.italic = True

# ── Persona card table ────────────────────────────────────────────────────────
card = doc.add_table(rows=2, cols=2)
card.style = "Table Grid"

LABEL_COLOR = "D9E2F3"

def card_label(cell, text):
    cell.text = text
    r = cell.paragraphs[0].runs[0]
    r.font.bold = True
    r.font.name = BODY_FONT
    r.font.size = Pt(9)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), LABEL_COLOR)
    tcPr.append(shd)

def card_value(cell, text):
    cell.text = text
    r = cell.paragraphs[0].runs[0]
    r.font.name = BODY_FONT
    r.font.size = Pt(10)

card_label(card.rows[0].cells[0], "BASIC PROFILE")
card_value(card.rows[0].cells[1],
    "Age: 21  |  Year: Third-year  |  Degree: Biomedical Engineering\n"
    "Location: On-campus student, university Wi-Fi\n"
    "Device: Personal Windows laptop (primary), phone for quick checks\n"
    "Tech comfort: Comfortable but not technical — uses Notion, Google Drive, "
    "Discord, Canvas daily; has never written a script or automation")

card_label(card.rows[1].cells[0], "USAGE CONTEXT")
card_value(card.rows[1].cells[1],
    "Frequency: 4–6 times per semester (semester start + each time a lecturer "
    "emails a date change)\n"
    "Choice vs. required: By choice — she actively looks for tools that reduce admin work\n"
    "Trigger: Opening a new syllabus PDF and realising it is 22 pages long")

doc.add_paragraph()

# ── Goals, Frustrations, Behaviors, Needs in a 4-row table ───────────────────
detail = doc.add_table(rows=4, cols=2)
detail.style = "Table Grid"

sections = [
    ("GOALS",
     "• Get every deadline from all five course syllabi into Google Calendar in "
     "one session at the start of term — not spread across five separate sittings.\n"
     "• Stop relying on study-group Discord pings as her primary deadline reminder system.\n"
     "• Catch mid-semester date changes before they cause a missed submission, not after.\n"
     "• Spend less than five minutes on calendar setup per course so she has time to "
     "actually read the material."),

    ("FRUSTRATIONS",
     "• Each syllabus is a 15–25 page PDF. Deadlines are scattered: one in the "
     "week-by-week schedule, another in the assessment overview table, a third buried "
     "in a late-penalty footnote. She has to read the whole document to be sure she "
     "found them all — and she almost never does.\n"
     "• She missed a lab report submission last semester because the due date was only "
     "mentioned in a footnote on page 19 of the course outline. She did not see it "
     "until the day after the deadline.\n"
     "• When lecturers send 'date change' emails mid-semester, she has to find the "
     "original calendar entry, remember which course it belongs to, and update it "
     "manually. She often forgets the update step.\n"
     "• Existing tools (Google Calendar, Notion, Canvas) do not read PDFs. She still "
     "has to do all the extraction by hand, which defeats the point of having a "
     "calendar app."),

    ("BEHAVIORS",
     "• Downloads all syllabi on the first day of semester but rarely reads them "
     "fully until the night before something is due.\n"
     "• Skims interfaces quickly — if important information is not visible in the "
     "first three seconds, she assumes it is not there.\n"
     "• Prefers browser-based tools over installing software ('I don't want "
     "another app on this laptop').\n"
     "• Will do a quick visual check of extracted results if the interface makes "
     "it feel fast — she is willing to spend 30 seconds scanning a list, not 5 minutes "
     "reading it.\n"
     "• Abandons tools that require account creation before showing her any value."),

    ("NEEDS & CONCERNS",
     "• Needs: A tool that accepts a PDF and returns a ready-to-import calendar "
     "file in under 30 seconds, with zero account setup.\n"
     "• Needs: Enough context shown next to each extracted date (title, source "
     "snippet) that she can spot a wrong entry at a glance without re-reading the PDF.\n"
     "• Concern: What if the tool misses a deadline? A false sense of completeness "
     "is worse than doing it manually, because she stops double-checking.\n"
     "• Concern: She does not know what Ollama or a language model is and does not "
     "need to — but she will lose trust immediately if the tool gives her one "
     "obviously wrong date."),
]

for i, (label, value) in enumerate(sections):
    card_label(detail.rows[i].cells[0], label)
    card_value(detail.rows[i].cells[1], value)

doc.add_paragraph()

# ── Connection to problem statement ──────────────────────────────────────────
body(doc, "Direct connection to the problem statement:", bold=True)
body(doc,
    "The problem statement identifies three specific pain points: deadlines buried across "
    "long PDFs, mid-semester email updates that are easy to miss, and the consequence of "
    "missing one date (a zero grade). Maya embodies all three. She receives 15–25 page "
    "syllabi she cannot fully read, she has already missed a deadline hidden in a footnote, "
    "and she has no reliable system for catching mid-semester date changes. Her technical "
    "comfort level is high enough to use a web tool without hand-holding, but low enough "
    "that she cannot build her own solution. She is exactly the person this tool is designed "
    "for — not because she is incompetent, but because the problem is genuinely tedious "
    "and the existing tools do not solve it."
)

# ═══════════════════════════════════════════════════════════════════════════════
#  3. SCENARIO
# ═══════════════════════════════════════════════════════════════════════════════
heading(doc, "Scenario")
body(doc,
    "It is the first Monday of semester. Maya has just received her four course outlines as "
    "PDF attachments. She opens the Deadline Extractor in her browser. The first thing she "
    "sees is the upload card. She drags her first syllabus PDF onto the PDF drop zone. "
    "The zone highlights briefly to confirm it received the file. She clicks the "
    "\"Extract Deadlines\" button."
)
body(doc,
    "The button changes to \"Extracting...\" while the system sends the file to the backend. "
    "The backend converts the PDF to markdown using pdfplumber, then sends that text to "
    "Gemma 4 e2b running locally through Ollama. About two to four seconds later, the panel "
    "switches to the results view. It shows nine deadlines, sorted by date."
)
body(doc,
    "Maya scans the list. Each row shows the date, an urgency badge (Overdue, Today, "
    "This Week, Upcoming), the title the model inferred, and a snippet of context from the "
    "original document. She notices one entry reads \"Date: 2026-11-20\" with no real title. "
    "She recognises that date as her final exam. The model missed the label. She makes a "
    "mental note. The other eight look right."
)
body(doc,
    "She clicks \"Export as .ics calendar file.\" A file named deadlines.ics downloads "
    "instantly. She opens it. Google Calendar asks if she wants to import 9 events. She "
    "clicks yes. All nine show up in her calendar with the correct dates. She goes back to "
    "the Deadline Extractor and repeats for her remaining three syllabi. The whole process "
    "takes under five minutes. She used to spend forty-five minutes doing this by hand "
    "at the start of every semester, if she did it at all."
)

# ═══════════════════════════════════════════════════════════════════════════════
#  4. HUMAN-AI INTERACTION DESIGN
# ═══════════════════════════════════════════════════════════════════════════════
heading(doc, "Human-AI Interaction Design")

body(doc, "What the user does:", bold=True)
bullet(doc, "Drops a PDF, uploads a screenshot, or pastes text into the upload panel.")
bullet(doc, "Clicks one button to trigger extraction.")
bullet(doc, "Reads through the list of extracted deadlines and filters by urgency if needed.")
bullet(doc, "Clicks \"Export as .ics calendar file\" to download a ready-to-import calendar file.")
bullet(doc, "Clicks \"← New input\" to start over with another document.")

body(doc, "What the model does:", bold=True)
body(doc,
    "Gemma 4 e2b receives the document content as text, or the image directly, through Ollama's "
    "chat API. Its only job is to read that content and return a JSON array. Each item in "
    "the array must have four fields: title (a short human-readable label), date (strictly "
    "YYYY-MM-DD), time (HH:MM in 24-hour format, defaulting to 23:59 if none is specified), "
    "and description (a sentence of context). The model is told explicitly in the system "
    "prompt: return only the JSON array, no explanation, no code fences, no commentary."
)

body(doc, "How the design handles model failures:", bold=True)
bullet(doc,
    "Code fence stripping: Despite the instruction, the model sometimes wraps its JSON in "
    "```json ... ``` fences. The backend strips these with a regex before attempting to parse, "
    "so the response still works."
)
bullet(doc,
    "JSON extraction fallback: If the response is not valid JSON at all, the backend runs a "
    "second regex looking for any [...] array anywhere in the string. If that also fails, "
    "it returns an empty list instead of crashing. The user sees zero deadlines and can retry."
)
bullet(doc,
    "Deterministic urgency sorting: The model is not asked to judge whether a deadline is "
    "overdue or upcoming. That calculation is done in the frontend with plain JavaScript "
    "date arithmetic. This avoids the model making wrong relative-time judgements like "
    "treating a past date as upcoming."
)
bullet(doc,
    "Mandatory human review: The extracted list is always shown to the user before export. "
    "The user can see the model's inferred title, the exact date, and a snippet of "
    "surrounding context. This gives the user a chance to catch hallucinations or missed "
    "labels before they land in their calendar."
)
bullet(doc,
    "Low temperature: The model is called with temperature = 0.1. This makes its output more "
    "consistent across repeated calls with the same input and reduces creative hallucination."
)
bullet(doc,
    "No relative date inference: If the model cannot determine a specific calendar date "
    "(for example, \"by end of term\" with no date anywhere nearby), the system prompt "
    "instructs it to omit the entry rather than guess. This reduces the chance of wrong "
    "dates appearing in the calendar."
)

# ═══════════════════════════════════════════════════════════════════════════════
#  5. MODEL LIMITS LOG
# ═══════════════════════════════════════════════════════════════════════════════
heading(doc, "Model Limits Log")
body(doc,
    "All rows are from real inputs tested during development. Each failure type is labelled "
    "with its code and a one-line explanation of the underlying mechanism — what is actually "
    "going wrong inside the model, not just what the symptom looks like."
)

headers = ["Input", "What failed", "Failure type & mechanism", "Mitigation tried", "Helped? (detail)"]
rows = [
    (
        '"Final project due end of term" — no date anywhere in the document',
        'Model returned {"date": "2026-01-31"} — a date it invented, not found in text.',
        "Hallucination (confabulation from implicit context): No date token exists to extract, "
        "so the model pattern-matches to what a 'final project' deadline typically looks like "
        "and generates a plausible-sounding date from prior training data instead of the document.",
        'System prompt updated: "If no specific calendar date can be determined from the text, '
        'omit the entry entirely."',
        "Yes — the fabricated date no longer appears. Trade-off: the entry is now silently "
        "dropped, so the user loses the reminder entirely. A real deadline with vague phrasing "
        "gets no output instead of a wrong one. Less dangerous, but still a gap.",
    ),
    (
        "Screenshot of a handwritten weekly timetable (photo quality, moderate contrast)",
        "Returned 3 dates not visible in the image; missed 2 dates that were clearly present.",
        "Hallucination (visual confabulation): The model generates expected-seeming content "
        "rather than grounding strictly to what it can read. It infers what a typical timetable "
        "should contain and fills gaps with invented dates.",
        "Human review step made mandatory — user always sees the list before exporting.",
        "Partially. The review step gives the user a chance to spot invented dates before they "
        "land in the calendar. But to confirm which entries are real, the user must "
        "cross-reference the original image — that takes time and partially defeats the "
        "purpose of automation.",
    ),
    (
        "32-page university course outline PDF containing 14 deadlines across all pages",
        "Returned 8 deadlines. The last 6 were silently dropped — no error, no truncation warning.",
        "Context-limit / forgetting (silent truncation): As the input token count grows, "
        "the model's effective attention over early content degrades. It does not signal "
        "when it stops processing — it simply returns fewer results as if they are complete.",
        "pdfplumber's structured markdown output is more compact than raw PDF text, buying "
        "some context headroom. No structural fix yet.",
        "No. The mitigation reduces how often this triggers (shorter compact input = fewer "
        "dropped items on medium-length docs) but does not prevent it on long documents. "
        "A 32-page syllabus still loses items. This is the most critical unresolved failure.",
    ),
    (
        '"Submit by the 15th" — day given, month and year missing',
        'Model returned {"date": "2026-01-15"} — January, which was already past.',
        "Hallucination + instruction-miss (partial date resolution): The model has one date "
        "token ('15th') and fills the missing month by anchoring to January, the most "
        "statistically common month in training examples with partial dates. It does not "
        "flag the ambiguity despite the instruction to omit unclear entries.",
        'Prompt updated: "Prefer the most logically upcoming date; if month cannot be '
        'determined at all, omit the entry."',
        "Partially. The model now more often picks a future month when one is inferable from "
        "surrounding context (e.g. a syllabus dated August will guess September). But if "
        "context is absent, it still picks January. The omit-if-ambiguous rule only fires "
        "when the model judges the entry ambiguous — which it often does not.",
    ),
    (
        "PDF with a 3-column assignment table: Assignment Name | Due Date | Weighting (%)",
        "Two adjacent table rows merged into one JSON entry. Title read 'Assignment 2 & Assignment 3'; date took the later of the two.",
        "Format-break (structure collapse): Without markdown conversion, pdfplumber returns "
        "table content as a flat string. Column boundaries disappear. The model sees "
        "'Assignment 2 15 Sep 20% Assignment 3 1 Oct 20%' as one run-on phrase and "
        "conflates adjacent rows into a single object.",
        "pdfplumber table extraction now converts each table to markdown (| col | col | col |) "
        "before sending. The model receives structural cues.",
        "Yes, for most tables. The markdown pipes restore enough visual structure for the model "
        "to treat each row as a separate entry. Edge case that remains: cells containing "
        "line breaks inside them still occasionally cause row misalignment.",
    ),
    (
        "Same 12-page syllabus uploaded twice in immediate succession (no change to input)",
        "Call 1 returned 9 deadlines. Call 2 returned 7 — a different subset, with 2 swapped and 3 dropped.",
        "Inconsistency (sampling variance compounding over long outputs): Even at moderate "
        "temperature, the autoregressive sampler chooses slightly different token paths on "
        "each call. For a short response this variance is unnoticeable. For a long JSON array "
        "(9 objects, hundreds of tokens), each small divergence early in the sequence "
        "compounds, producing a meaningfully different end result.",
        "Temperature reduced from default to 0.1 — the lowest stable value before the model "
        "starts looping or truncating.",
        "Mostly yes. At 0.1, repeated calls on the same input now return the same result "
        "in roughly 90% of tests. The remaining 10% variance is concentrated on borderline "
        "entries (dates with ambiguous wording) — exactly the ones a human reviewer should "
        "double-check anyway.",
    ),
    (
        '"Due: 07/08/2026" — date string interpretable as either July 8 (MM/DD) or August 7 (DD/MM)',
        "Model silently chose MM/DD (July 8) with no flag, no note in the description field.",
        "Instruction-miss (format ambiguity not surfaced): The system prompt specifies "
        "YYYY-MM-DD output but gives no instruction on how to handle ambiguous input formats. "
        "The model defaults to MM/DD (US convention) because that format dominates its "
        "training data, and it does not recognise the ambiguity as worth flagging.",
        "Prompt now notes: 'If a date format is ambiguous, note the ambiguity in the "
        "description field.' Output format requirement (YYYY-MM-DD) unchanged.",
        "Partially. The model now sometimes notes ambiguity in the description (e.g. "
        "'Note: date format unclear — assumed MM/DD'). But it does so inconsistently — "
        "roughly 50% of the time. The core problem (model picks one interpretation and "
        "commits) remains. Full fix would require the user to specify their locale, "
        "which adds friction we have not built yet.",
    ),
    (
        "Image of a printed exam timetable — small font (~8pt equivalent), two-column layout",
        "Model returned an empty array. No error, no partial result, no explanation.",
        "Refusal / over-caution (confidence threshold silence): Rather than attempting "
        "extraction with caveats, the model returns nothing when its internal confidence "
        "in reading the image falls below a threshold. This manifests identically to "
        "a genuine zero-deadline result, making it impossible to distinguish 'nothing found' "
        "from 'gave up silently.'",
        "System prompt updated: 'Extract all dates you can identify. If the image is low "
        "quality, include your best attempt and note uncertainty in the description field. '",
        "Yes. The model now attempts extraction even on difficult images and returns "
        "whatever it can read, flagging uncertain entries with notes like 'low image "
        "quality — verify this date.' Accuracy on poor-quality images remains low, but the "
        "user now gets something to review rather than a blank page.",
    ),
]
add_table(doc, headers, rows)

# ═══════════════════════════════════════════════════════════════════════════════
#  6. REFLECTIONS
# ═══════════════════════════════════════════════════════════════════════════════
heading(doc, "Reflections")

heading(doc, "1. User Needs", level=2)
body(doc,
    "Maya's primary need is time savings with acceptable accuracy. She does not need a "
    "perfect tool. She needs a tool that is fast, that gives her most of the deadlines "
    "correctly, and that lets her quickly catch the ones it gets wrong before they "
    "matter. Secondary needs are trust (she has to believe the output is worth checking) "
    "and zero setup friction (no account, no installation, no config)."
)

heading(doc, "2. Why This Design Meets Those Needs", level=2)
body(doc,
    "The pipeline architecture keeps the AI responsible for only the genuinely hard part: "
    "understanding natural language and extracting structured date information from "
    "unstructured text. Everything else in the pipeline is deterministic and reliable. "
    "pdfplumber does structured PDF-to-text conversion with no AI involved. The frontend "
    "does urgency sorting with plain arithmetic. The icalendar library generates the .ics "
    "file from a Python dictionary with no AI involved. This means the system's failure "
    "surface is limited mostly to the LLM step, and the human review panel is placed "
    "exactly where the risk is highest, right after that step."
)
body(doc,
    "The single-button interaction matches Maya's expectation that things work immediately. "
    "There is no configuration screen, no settings, no account. Upload, click, review, export. "
    "The urgency colour-coding means Maya can glance at the list and immediately see if "
    "anything is this week without reading every row."
)

heading(doc, "3. How Well It Works", level=2)
body(doc,
    "For clean, well-formatted text, the system performs well. On test runs with typical "
    "university syllabi, it correctly extracts roughly 85 to 90 percent of deadlines on the "
    "first try, with most failures being missed dates rather than wrong dates. Wrong dates "
    "(hallucinations) are rarer but more dangerous, which is exactly why the review step "
    "exists."
)
body(doc,
    "Performance degrades significantly for very long documents (over 25 pages), scanned or "
    "handwritten images with small text, and documents with ambiguous date formats or "
    "relative language. These are real limits, not edge cases, for the student use case."
)

heading(doc, "4. What Else Could Work", level=2)
bullet(doc,
    "Rule-based regex extraction: Fast, zero hallucination, zero cost. But it misses "
    "natural-language dates like 'two weeks from submission' and requires manual format "
    "rules for every date style."
)
bullet(doc,
    "A fine-tuned model dedicated to date extraction: More reliable than a general-purpose "
    "chat model for this narrow task. Requires a labelled training dataset of documents "
    "with ground-truth date annotations."
)
bullet(doc,
    "A retrieval-augmented approach: Pre-scan the document with regex to find candidate "
    "date strings, then send only those strings plus their surrounding context to the LLM. "
    "This would reduce the model's input length and focus it on the hard cases only."
)
bullet(doc,
    "Browser extension: Instead of a separate app, inject the extraction tool directly into "
    "the PDF viewer or email client. Lower friction for users who already have the document open."
)

heading(doc, "5. What Gemma 4's Limits Forced the Design to Work Around", level=2)
body(doc,
    "The most consequential constraint was the context window. Gemma 4 e2b handles moderate "
    "length documents well but starts dropping entries on longer inputs. This is not a "
    "graceful failure. The model does not say 'I stopped at page 10.' It just silently "
    "returns fewer results. This forced the inclusion of pdfplumber's structured extraction, "
    "which produces more compact markdown than raw PDF text, buying extra context headroom. "
    "It also forced an honest note in the design: the system cannot guarantee completeness "
    "on very long documents, and users should check their results."
)
body(doc,
    "The model's inconsistency at higher temperatures forced a design choice: run at "
    "temperature 0.1 and accept slightly more repetitive outputs in exchange for stable "
    "results. A user who reloads the page and re-uploads the same PDF should not get a "
    "different set of deadlines."
)
body(doc,
    "The model's tendency to ignore output format instructions forced two layers of "
    "defensive parsing in the backend: a code-fence stripper and a fallback JSON array "
    "extractor. Without these, a single model response with an extra sentence before the "
    "JSON would crash the entire extraction."
)
body(doc,
    "The model's hallucination on ambiguous inputs forced the design to prioritise the "
    "human review step as a non-optional part of the flow. The results panel cannot be "
    "skipped. The user always sees the list before downloading. This was not the original "
    "design, which had an option to auto-download. That was removed specifically because "
    "the model is not reliable enough to skip human eyes."
)

# ═══════════════════════════════════════════════════════════════════════════════
#  6b. REPORT CONSTRAINTS & WHAT COULD BE IMPROVED
# ═══════════════════════════════════════════════════════════════════════════════
heading(doc, "Report Constraints & What Could Be Improved")
body(doc,
    "This report has real limits that should be named honestly rather than papered over."
)

body(doc, "Constraints on the current report:", bold=True)
bullet(doc,
    "No formal evaluation data. The 85-90% accuracy figure quoted in Reflections is an "
    "informal estimate from watching test runs, not a measured precision/recall score "
    "against a labelled ground-truth set. It could be significantly off in either direction."
)
bullet(doc,
    "Persona is hypothetical. Maya Chen was constructed from general knowledge of student "
    "behaviour, not from user interviews or surveys. Her actual needs and frustrations may "
    "differ from what is described here."
)
bullet(doc,
    "Small failure sample. The model limits log covers 8 failure types observed during "
    "development. A wider test across more document types, languages, and formats would "
    "likely reveal additional failure modes not captured here."
)
bullet(doc,
    "Binary helped/not-helped framing understates nuance. Even where a mitigation is marked "
    "as working, it often introduced a trade-off (e.g. fixing hallucination by omitting "
    "entries also drops legitimate vague deadlines). The log now describes these trade-offs "
    "but a fuller account would quantify them."
)
bullet(doc,
    "No latency data. Response time varies by document length and machine load but no "
    "systematic measurements were taken. The claim that extraction takes 'two to four seconds' "
    "is observational and not benchmarked."
)

body(doc, "What could be improved in the report:", bold=True)
bullet(doc,
    "Run a benchmark. Test 15-20 real syllabi with manually annotated dates and report "
    "actual precision (how many extracted dates were correct) and recall (how many real "
    "dates were found). This replaces the informal estimate with a defensible number."
)
bullet(doc,
    "Validate the persona through at least three short interviews with real students. "
    "Check whether the stated needs, friction points, and willingness to do a review step "
    "actually hold."
)
bullet(doc,
    "Expand the limits log with longer-tail failures: non-English documents, PDFs with "
    "scanned image pages (no text layer), heavily formatted academic papers with footnotes "
    "containing dates, and recurring events (every Friday) that the model should either "
    "expand or flag."
)
bullet(doc,
    "Add a scenario where the tool fails visibly and describe what the user experience looks "
    "like. The current scenario only covers the happy path. A failure scenario would show "
    "how the design degrades gracefully rather than just how it works when everything goes right."
)

# ═══════════════════════════════════════════════════════════════════════════════
#  7. PLAN
# ═══════════════════════════════════════════════════════════════════════════════
heading(doc, "Plan")

heading(doc, "Mitigations for Model Limits", level=2)
bullet(doc,
    "Context window: Split long PDFs into page-sized chunks and run extraction on each chunk "
    "independently. Merge the results with deduplication (same date + similar title = one "
    "entry). This keeps each LLM call short and avoids the silent-dropping problem."
)
bullet(doc,
    "Inconsistency: Run the model twice on high-stakes documents and take the union of both "
    "result sets after deduplication. Show the user a confidence indicator (one result from "
    "both calls = high confidence, only one call = lower)."
)
bullet(doc,
    "Hallucination on ambiguous dates: Add a regex pre-pass that identifies candidate date "
    "strings in the raw text. Only send the model those strings plus 80 characters of "
    "surrounding context. This narrows the task and gives the model much less room to invent."
)
bullet(doc,
    "Format-break on structured content: Already partly addressed with pdfplumber's table "
    "extraction. Next step: provide the model with examples of table rows in the few-shot "
    "section of the system prompt so it learns the expected structure from context."
)

heading(doc, "AI + Non-AI Methods", level=2)
body(doc,
    "The current pipeline already combines both. pdfplumber (non-AI) handles the PDF layer. "
    "Gemma 4 e2b (AI) handles the language layer. icalendar (non-AI) handles the output layer. "
    "The plan extends this:"
)
bullet(doc, "Add regex pre-screening (non-AI) as a first pass to anchor the model's task.")
bullet(doc, "Add a confidence score display based on how many model runs agreed on each deadline.")
bullet(doc, "Keep the export pipeline fully deterministic. The AI never touches the .ics file directly.")

heading(doc, "What to Evaluate Next", level=2)
bullet(doc,
    "Precision and recall on a benchmark set of 15 real course syllabi with manually "
    "annotated ground-truth deadlines. This gives a concrete accuracy number."
)
bullet(doc,
    "User study with 5-8 students. Measure: time from opening the tool to all deadlines in "
    "calendar, number of errors the review step caught, number of errors that slipped through "
    "into the calendar, and subjective confidence in the tool."
)
bullet(doc,
    "Track how often the human review step actually leads to a user spotting and discarding "
    "a deadline. This tells us the real error rate in production and whether the review step "
    "is worth the friction it adds."
)
bullet(doc,
    "Latency benchmarks across document sizes. Identify the document length at which "
    "response time becomes unacceptable to users (hypothesis: over 10 seconds)."
)

# ═══════════════════════════════════════════════════════════════════════════════
#  8. AI-USE ACKNOWLEDGMENT
# ═══════════════════════════════════════════════════════════════════════════════
heading(doc, "AI-Use Acknowledgment")
body(doc,
    "The following AI tools were used in this project. All use is disclosed in accordance "
    "with the course plagiarism policy."
)

body(doc, "Gemma 4 e2b (Google, accessed via Ollama):", bold=True)
body(doc,
    "This is the primary AI model in the system. It runs at inference time to extract "
    "deadline information from document content. It is not used during development, report "
    "writing, or any other part of the project outside of its designated role in the "
    "extraction pipeline."
)

body(doc, "Claude (Anthropic, accessed via Claude Code):", bold=True)
body(doc,
    "Claude was used as a coding assistant throughout the development of this project. "
    "Specifically, it helped write the FastAPI backend (main.py, deadline_extractor.py, "
    "pdf_processor.py, ics_generator.py), the React/TypeScript frontend components "
    "(UploadPanel, ResultsPanel, DeadlineRow, App), and the Vite configuration. Claude "
    "also helped debug integration issues between the frontend and backend, including the "
    "proxy setup and the multipart form data handling."
)
body(doc,
    "Claude was additionally used to help write and structure this report. The core ideas, "
    "the persona, the scenario, the design decisions, and the observed model failures all "
    "came from the author's own work building and testing the system. Claude helped turn "
    "those ideas into coherent written paragraphs and helped structure the report to meet "
    "the assignment requirements. The content is the author's own; the writing assistance "
    "is Claude's."
)
body(doc,
    "No AI tool was used to fabricate observations, invent failure examples, or generate "
    "content that misrepresents the author's own understanding of the system. The model "
    "limits log reflects real failures encountered during development and testing, not "
    "hypothetical ones invented for the report."
)

# ── Save ──────────────────────────────────────────────────────────────────────
doc.save(OUT)
print(f"Saved: {OUT}")
